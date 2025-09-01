import streamlit as st
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")  # ✅ Safe backend for Streamlit
import matplotlib.pyplot as plt
import seaborn as sns
from lifelines import KaplanMeierFitter, CoxPHFitter
from lifelines.statistics import logrank_test, multivariate_logrank_test
import base64, time, os
from io import BytesIO
from docx import Document

# === Styling ===
st.markdown("""
    <style>
    body {
        background: linear-gradient(to bottom, #e6f0ff, #ffffff);
    }
    [data-testid="stAppViewContainer"] {
        background: linear-gradient(to bottom, #e6f0ff, #ffffff);
    }
    [data-testid="stSidebar"] > div:first-child {
        background-color: #d1e4f9 !important;
        border-radius: 16px;
        padding: 1rem;
    }
    .stButton > button, .stSelectbox div, .stTextInput > div > input {
        border-radius: 12px !important;
    }
    </style>
""", unsafe_allow_html=True)

# === Splash Screen ===
def get_base64_image(image_path):
    with open(image_path, "rb") as f:
        return base64.b64encode(f.read()).decode()

# ✅ Use a single correct logo path
current_dir = os.path.dirname(os.path.abspath(__file__))
logo_path = os.path.join(current_dir, "assets", "logo.png")

if "splash_done" not in st.session_state:
    if os.path.exists(logo_path):
        img_base64 = get_base64_image(logo_path)
        st.markdown(f"""
            <style>
            .splash {{
                position: fixed; top: 0; left: 0; right: 0; bottom: 0;
                background-color: black;
                display: flex; justify-content: center; align-items: center;
                z-index: 9999;
            }}
            .splash img {{ width: 100vw; height: 100vh; object-fit: cover; }}
            </style>
            <div class="splash"><img src="data:image/png;base64,{img_base64}" /></div>
        """, unsafe_allow_html=True)
        time.sleep(5)
        st.session_state.splash_done = True
        st.rerun()
# === App Title ===
st.markdown(f"""
    <div style='text-align:center;'>
        <img src='data:image/png;base64,{get_base64_image("assets/logo.png")}' style='height:80px;'>
        <h1>OncoClini Partner</h1>
    </div>
""", unsafe_allow_html=True)

# === Navigation ===
section = st.radio("Select View", ["Home", "Analysis", "Library", "Tutorial"], horizontal=True, key="section_radio")

# === Home Section ===
if section == "Home":
    st.info("Please select a section to proceed.")

    st.markdown("""
        <style>
        .home-card {
            background: #f5faff;
            padding: 24px;
            border-radius: 20px;
            box-shadow: 0 8px 20px rgba(0, 0, 0, 0.08);
            margin-bottom: 25px;
        }
        .home-title {
            font-size: 22px;
            font-weight: bold;
            margin-bottom: 6px;
            display: flex;
            align-items: center;
        }
        .home-icon {
            font-size: 26px;
            margin-right: 10px;
        }
        .home-text {
            font-size: 16px;
            margin-left: 35px;
            color: #333;
        }
        </style>

        <div class="home-card">
            <div class="home-title">
                <span class="home-icon">🔍</span> What is OncoClini Partner?
            </div>
            <div class="home-text">
                OncoClini Partner is a no-code survival analysis tool tailored for oncology professionals.
                It simplifies Kaplan-Meier and Cox models into one-click insights — making analysis faster, accurate, and accessible.
            </div>
        </div>

        <div class="home-card">
            <div class="home-title">
                <span class="home-icon">👥</span> Who Can Use It?
            </div>
            <div class="home-text">
                This tool is ideal for:
                <ul>
                    <li>👨‍⚕️ Doctors & Oncologists</li>
                    <li>🔬 Clinical Researchers</li>
                    <li>🏢 CROs / CDMOs</li>
                </ul>
                No programming skills required — just your dataset and a question.
            </div>
        </div>

        <div class="home-card">
            <div class="home-title">
                <span class="home-icon">🎯</span> Our Vision
            </div>
            <div class="home-text">
                To empower oncology professionals with automated statistical tools that simplify time-to-event analysis, enhance patient care, and improve outcomes.
            </div>
        </div>
    """, unsafe_allow_html=True)

# === Library Section ===
elif section == "Library":
    st.subheader("📘 Survival Analysis – FAQ")

    content_text = """
1. **What is Survival Analysis? Why Do We Use It?**

Survival analysis is a statistical method used to study time until an event happens, such as death, relapse, or recovery.

**Why we use it:**
- It can handle patients who are still alive (censored data)
- It deals with different follow-up times for different patients
- It estimates survival probability at different time points
- It helps compare treatments based on how long patients live

**Why not use regular regression (like linear or logistic):**

- Linear regression assumes continuous outcome (time), but cannot handle censoring
- Logistic regression only gives yes/no outcome, not how long it took
- Survival analysis focuses on "when" an event happens, not just "whether" it happens

---
2. **What is a Kaplan-Meier (KM) Curve?**

A KM curve is a graph showing the percentage of patients surviving over time.
It steps down each time a patient dies or relapses.

**How to read it:**
- X-axis = Time (months or years)
- Y-axis = % of patients alive or event-free
- Each step down = One or more patients had the event
- Flat parts = No event happened during that time

**Example:**
 At 12 months, if the curve shows 70%, it means 70% of patients are still alive.

**Assumption:**
- All patients have the same probability of being followed
- Events occur independently

**Limitation:**
- Cannot adjust for multiple factors (e.g., age, stage)
- Cannot model continuous covariates

---
3. **What is Censoring?**

Censoring means the patient's exact event time is unknown.

**This can happen when:**
- The patient is still alive when the study ends
- The patient drops out or is lost to follow-up

**Example:**
 If a patient is alive at the end of the 10-month study, we record 10 months and censor them.

**Assumption:**
- Censored patients are assumed to have the same risk as others still in the study

**Limitation:**
- Too many censored patients can weaken the study conclusions

---
4. **What is the Log-Rank Test?**

This is a statistical test used to compare survival between two or more groups (e.g., two treatments).
It checks whether the difference between curves is real or due to chance.

**Example:**
 Comparing survival between Treatment A and Treatment B.
If p-value = 0.02, the difference is likely real.

**Assumption:**
- Hazard (risk) is proportional across time
- Censoring is non-informative

**Limitation:**
- Cannot adjust for multiple factors
- Only valid if hazards stay constant over time

---
5. **What is a Hazard Ratio (HR)?**

Hazard Ratio compares the risk of event in one group to another over time.

- HR = 1 → Both groups have same risk
- HR < 1 → Treatment group has lower risk
- HR > 1 → Treatment group has higher risk

**Example:**
 HR = 0.70 means 30% lower risk of death in the treatment group.

**Assumption:**
- Risk ratio (hazard) between groups is constant over time

**Limitation:**
- If the hazard changes over time, the result may not be accurate

---
6. **What is Confidence Interval (CI) in HR?**

A confidence interval shows the range where the true HR is likely to fall.

- If the CI does not include 1, the result is statistically significant
- A narrow CI means the estimate is more reliable

**Example:**
 HR = 0.65 (95% CI: 0.50 – 0.85)
This means we are 95% confident the true HR is between 0.50 and 0.85.

**Limitation:**
- If the CI is too wide, the result is less reliable

---
7. **What is Median Survival Time?**

It is the time when half of the patients have had the event (death or relapse), and half are still event-free.

**Example:**
 If 5 out of 10 patients die by 12 months, median survival is 12 months.

**Assumption:**
- Event and censoring times are accurately recorded

**Limitation:**
- Doesn’t show the full survival experience (just 50% point)

---
8. **What is a Cox Proportional Hazards Model?**

This is a method to find out which factors affect survival (like age, treatment, stage).
It adjusts for multiple variables at once.

**Example:**
 It may show that older age and advanced cancer stage are linked to shorter survival.

**Assumption:**
- Hazard ratios are proportional (constant) over time
- The effects of variables are additive in log-scale

**Limitation:**
- If hazard rates change over time, results may be misleading
- Needs a larger sample size to adjust for multiple factors

---
9. **What is an Event in Survival Analysis?**

An event is the outcome we are studying, such as death, relapse, or disease progression.

**Example:**
 In a lung cancer study, the event may be death from any cause.

**Assumption:**
- Events are correctly defined and consistently recorded

---
10. **How to Read KM Curves on the Dashboard?**
- X-axis: Time (e.g., months or years)
- Y-axis: % of patients who are alive or event-free
- Sharp drops: Sudden events
- Flat parts: No events
- Compare different groups by different colored lines

**Example:**
If one curve stays higher, it means better survival in that group.

---
11. **How Can Doctors Use This Dashboard in Practice?**
- Compare treatments to see which has better survival
- Identify which patient groups are at higher risk
- Support patient communication and shared decision-making
- Understand time-to-event, not just event occurrence

**Limitation:**
- This is a decision support tool, not a final treatment guide
- Always combine with clinical experience and judgment


    """

    st.markdown(content_text)


# === Tutorial Section ===
elif section == "Tutorial":
    st.subheader("📽️ Tutorial")
    st.video("assets/survival_tutorial.mp4")
elif section == "Analysis":
    st.sidebar.markdown("### 📁 Upload your dataset")
    uploaded_file = st.sidebar.file_uploader("Upload file", type=["csv", "xlsx", "xls", "json"])

    if uploaded_file:
        ext = uploaded_file.name.split(".")[-1].lower()
        if ext == "csv":
            df = pd.read_csv(uploaded_file)
        elif ext in ["xlsx", "xls"]:
            df = pd.read_excel(uploaded_file)
        elif ext == "json":
            df = pd.read_json(uploaded_file)
        else:
            st.error("Unsupported file format"); st.stop()

        file_prefix = os.path.splitext(uploaded_file.name)[0]
        st.session_state.df = df
        st.session_state.file_prefix = file_prefix

        st.subheader("📊 Uploaded Data Preview")
        st.dataframe(df.head(), use_container_width=True)
        st.session_state.data_preview = df.head()

        # === Select Analysis
        analysis = st.sidebar.selectbox("Select Analysis", [
            "-- Select --",
            "Descriptive Stats",
            "Kaplan-Meier + Log-Rank Test",
            "Cox Proportional Hazards"
        ])
        st.session_state.selected_analysis = analysis
        st.session_state.setdefault("run_desc", False)
        st.session_state.setdefault("run_km", False)
        st.session_state.setdefault("run_cox", False)

        # === Descriptive Stats
        if analysis == "Descriptive Stats":
            st.session_state.run_desc = True
            cat_cols = df.select_dtypes(include='object').columns
            num_cols = df.select_dtypes(include=np.number).columns

            if len(cat_cols):
                st.subheader("🧾 Categorical Summary")
                cat_df = pd.DataFrame([
                    {
                        "Variable": col,
                        "Unique": df[col].nunique(),
                        "Missing": df[col].isnull().sum()
                    }
                    for col in cat_cols
                ])
                st.session_state.cat_table = cat_df
                st.dataframe(cat_df, use_container_width=True)

            if len(num_cols):
                st.subheader("📈 Numerical Summary")
                num_df = []
                for col in num_cols:
                    desc = df[col].describe()
                    iqr = desc['75%'] - desc['25%']
                    outliers = df[(df[col] < desc['25%'] - 1.5 * iqr) | (df[col] > desc['75%'] + 1.5 * iqr)].shape[0]
                    num_df.append({
                        "Variable": col,
                        "Mean": round(desc['mean'], 2),
                        "Median": round(df[col].median(), 2),
                        "Std": round(desc['std'], 2),
                        "Min": round(desc['min'], 2),
                        "Max": round(desc['max'], 2),
                        "Outliers": outliers
                    })
                num_df = pd.DataFrame(num_df)
                st.session_state.num_table = num_df
                st.dataframe(num_df, use_container_width=True)        
# === Kaplan-Meier + Log-Rank Test ===
        if analysis == "Kaplan-Meier + Log-Rank Test":
            st.session_state.run_km = True
            st.subheader("📉 Kaplan-Meier + Log-Rank Test")

            duration = st.selectbox("Duration Column", df.columns)
            event = st.selectbox("Event Column", df.columns)
            group = st.selectbox("Group Column", ["None"] + list(df.columns))

            kmf = KaplanMeierFitter()
            fig_km, ax = plt.subplots()

            if group == "None":
                kmf.fit(df[duration], event_observed=df[event])
                kmf.plot_survival_function(ax=ax)
            else:
                unique_groups = df[group].dropna().unique()
                for g in unique_groups:
                    subset = df[df[group] == g]
                    kmf.fit(subset[duration], event_observed=subset[event], label=str(g))
                    kmf.plot_survival_function(ax=ax)

                if len(unique_groups) == 2:
                    g1, g2 = unique_groups[:2]
                    d1 = df[df[group] == g1]
                    d2 = df[df[group] == g2]
                    result = logrank_test(d1[duration], d2[duration], d1[event], d2[event])
                    st.session_state.last_pvalue = result.p_value
                    st.markdown(f"**Log-Rank p-value (2 groups):** {result.p_value:.4f}")
                elif len(unique_groups) > 2:
                    result = multivariate_logrank_test(df[duration], df[event], df[group])
                    st.session_state.last_pvalue = result.p_value
                    st.markdown(f"**Log-Rank p-value (multi-group):** {result.p_value:.4f}")
                    st.markdown(f"Chi-square: {result.test_statistic:.2f} (df = {len(unique_groups) - 1})")

            st.pyplot(fig_km)
            st.session_state.km_plot = fig_km       
 # === Cox Proportional Hazards ===
        if analysis == "Cox Proportional Hazards":
            st.session_state.run_cox = True
            st.subheader("📌 Cox Proportional Hazards Model")

            duration = st.selectbox("Duration", df.columns)
            event = st.selectbox("Event", df.columns)
            covs = st.multiselect("Covariates", [c for c in df.columns if df[c].dtype != 'object' and c not in [duration, event]])

            if covs:
                cph = CoxPHFitter()
                data = df[[duration, event] + covs].dropna()
                cph.fit(data, duration_col=duration, event_col=event)

                st.write(cph.summary)
                st.session_state.cox_plot = cph.plot().figure
                st.pyplot(st.session_state.cox_plot)

                # === Interpretation ===
                st.subheader("🧠 Cox Model Interpretation")
                interpretation_text = []
                for index, row in cph.summary.iterrows():
                    hr = row['exp(coef)']
                    lower = row['exp(coef) lower 95%']
                    upper = row['exp(coef) upper 95%']
                    ci_text = f"(95% CI: {lower:.3f} to {upper:.3f})"
                    conclusion = "\u2705 Significant" if lower > 1 or upper < 1 else "\u26A0\uFE0F Not significant"
                    effect = "higher hazard" if hr > 1 else ("lower hazard" if hr < 1 else "no effect")
                    text = f"- **{index}**: HR = {hr:.3f} {ci_text} → {effect}, {conclusion}"
                    st.markdown(text)
                    interpretation_text.append(text)

                st.session_state.cox_interp = interpretation_text        
# === Chart Options ===
        chart_options = []
        if len(df.select_dtypes(include='object').columns):
            chart_options.extend(["Bar Chart", "Pie Chart"])
        if len(df.select_dtypes(include='number').columns):
            chart_options.extend(["Histogram", "Box Plot"])
        if len(df.select_dtypes(include='number').columns) >= 2:
            chart_options.append("Scatter Plot")

        chart_types = st.sidebar.multiselect("Select Chart Types", chart_options)
        st.session_state.chart_types = chart_types

        chart_figs = []
        if chart_types:
            st.subheader("📊 Selected Charts")
            for chart in chart_types:
                fig, ax = plt.subplots()
                st.markdown(f"#### {chart}")
                
                if chart == "Histogram":
                    hist_col = st.selectbox("Select variable for Histogram", df.select_dtypes(include='number').columns, key=f"{chart}_hist")
                    df[hist_col].plot(kind='hist', ax=ax)
                    ax.set_xlabel(hist_col)
                    ax.set_ylabel("Frequency")

                elif chart == "Box Plot":
                    sns.boxplot(data=df.select_dtypes(include='number'), ax=ax)
                    ax.set_title("Box Plot of Numeric Variables")

                elif chart == "Scatter Plot":
                    numeric = df.select_dtypes(include='number')
                    if numeric.shape[1] >= 2:
                        x_var = st.selectbox("X-axis", numeric.columns, key="scatter_x")
                        y_var = st.selectbox("Y-axis", numeric.columns, key="scatter_y")
                        ax.scatter(df[x_var], df[y_var])
                        ax.set_xlabel(x_var)
                        ax.set_ylabel(y_var)
                        ax.set_title(f"Scatter Plot: {x_var} vs {y_var}")

                elif chart == "Bar Chart":
                    col = st.selectbox("Bar Chart Variable", df.select_dtypes(include='object').columns, key="bar")
                    df[col].value_counts().plot(kind='bar', ax=ax)
                    ax.set_title(f"Bar Chart: {col}")
                    ax.set_ylabel("Count")

                elif chart == "Pie Chart":
                    col = st.selectbox("Pie Chart Variable", df.select_dtypes(include='object').columns, key="pie")
                    df[col].value_counts().plot(kind='pie', ax=ax, autopct='%1.1f%%')
                    ax.set_ylabel("")
                    ax.set_title(f"Pie Chart: {col}")

                st.pyplot(fig)
                chart_figs.append((chart, fig))

        st.session_state.chart_figs = chart_figs         
# === Chatbot Sidebar ===
        st.sidebar.markdown("### 🤖 Ask KASSR (Chatbot)")
        q = st.sidebar.text_input("Ask a question:")

        if "chat_log" not in st.session_state:
            st.session_state.chat_log = []

        if q:
            q_lower = q.lower()
            answered = False
            response = ""

            # === Cox p-value based query ===
            if "cox" in q_lower and "p-value" in q_lower:
                if "cox_interp" in st.session_state:
                    for line in st.session_state.cox_interp:
                        if "p-value" in line:
                            response = line
                            answered = True

            # === Kaplan-Meier Log-Rank Test p-value ===
            elif "p-value" in q_lower or "log-rank" in q_lower:
                if "last_pvalue" in st.session_state:
                    p = st.session_state["last_pvalue"]
                    result = "Significant" if p < 0.05 else "Not Significant"
                    response = f"Log-Rank p-value: {p:.4f} → {result}"
                    answered = True
                else:
                    response = "Run Kaplan-Meier first."
                    answered = True

            # === Mean or Median of a specific variable ===
            elif "mean" in q_lower or "median" in q_lower:
                found = False
                if "num_table" in st.session_state:
                    for _, row in st.session_state.num_table.iterrows():
                        if row['Variable'].lower() in q_lower:
                            mean = row['Mean']
                            median = row['Median']
                            response = f"The mean of {row['Variable']} is {mean}, and median is {median}."
                            found = True
                            answered = True
                if not found:
                    response = "Variable not found in numeric summary."
                    answered = True

            # === List all means ===
            elif "all means" in q_lower or "list means" in q_lower:
                if "num_table" in st.session_state:
                    lines = [f"{row['Variable']}: {row['Mean']}" for _, row in st.session_state.num_table.iterrows()]
                    response = "\n".join(lines)
                    answered = True

            # === Default fallback ===
            if not answered:
                response = "I couldn't understand. Try asking about mean, median, p-value etc."

            st.sidebar.info(response)
            st.session_state.chat_log.append((q, response))         
        

        def generate_doc():
            doc = Document()
            doc.add_heading("OncoClini Partner Report", 0)

            # === Data Preview ===
            if "data_preview" in st.session_state:
                doc.add_heading("Data Preview", level=1)
                df_preview = st.session_state["data_preview"]
                table = doc.add_table(rows=1, cols=len(df_preview.columns))
                for i, col in enumerate(df_preview.columns):
                    table.cell(0, i).text = col
                for _, row in df_preview.iterrows():
                    cells = table.add_row().cells
                    for i, val in enumerate(row):
                        cells[i].text = str(val)

            # === Descriptive Statistics ===
            if st.session_state.get("run_desc"):
                doc.add_heading("Descriptive Statistics", level=1)

                # Categorical Summary
                if "cat_table" in st.session_state:
                    doc.add_heading("Categorical Summary", level=2)
                    table = doc.add_table(rows=1, cols=len(st.session_state.cat_table.columns))
                    for i, col in enumerate(st.session_state.cat_table.columns):
                        table.cell(0, i).text = col
                    for _, row in st.session_state.cat_table.iterrows():
                        cells = table.add_row().cells
                        for i, val in enumerate(row):
                            cells[i].text = str(val)

                # Numerical Summary
                if "num_table" in st.session_state:
                    doc.add_heading("Numerical Summary", level=2)
                    table = doc.add_table(rows=1, cols=len(st.session_state.num_table.columns))
                    for i, col in enumerate(st.session_state.num_table.columns):
                        table.cell(0, i).text = col
                    for _, row in st.session_state.num_table.iterrows():
                        cells = table.add_row().cells
                        for i, val in enumerate(row):
                            cells[i].text = str(val)

            # === KM Plot & Log-Rank Test ===
            if st.session_state.get("run_km"):
                doc.add_heading("Kaplan-Meier Survival Analysis", level=1)
                if "last_pvalue" in st.session_state:
                    doc.add_paragraph(f"Log-Rank Test p-value: {st.session_state['last_pvalue']:.4f}")
                if "km_plot" in st.session_state:
                    img_path = "km_plot.png"
                    st.session_state.km_plot.savefig(img_path)
                    doc.add_picture(img_path)
                    os.remove(img_path)

            # === Cox Model ===
            if st.session_state.get("run_cox") and "cox_interp" in st.session_state:
                doc.add_heading("Cox Model Interpretation", level=1)
                for line in st.session_state.cox_interp:
                    doc.add_paragraph(line)
                if "cox_plot" in st.session_state:
                    cox_img_path = "cox_plot.png"
                    st.session_state.cox_plot.savefig(cox_img_path)
                    doc.add_picture(cox_img_path)
                    os.remove(cox_img_path)

            # === Charts ===
            if "chart_figs" in st.session_state and st.session_state.chart_figs:
                doc.add_heading("Charts", level=1)
                for i, (title, fig) in enumerate(st.session_state.chart_figs):
                    chart_path = f"chart_{i}.png"
                    fig.savefig(chart_path)
                    doc.add_paragraph(title)
                    doc.add_picture(chart_path)
                    os.remove(chart_path)

            # === Optional Chatbot Conversation ===
            if st.session_state.get("chat_log"):
                doc.add_heading("Chatbot Q&A", level=1)
                for q, a in st.session_state.chat_log:
                    doc.add_paragraph(f"You: {q}")
                    doc.add_paragraph(f"Bot: {a}")

            buffer = BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

        if st.session_state.selected_analysis != "-- Select --":
            st.sidebar.markdown("### 📄 Export")
            if st.sidebar.button("📥 Download Report"):
                word_data = generate_doc()
                name = f"{st.session_state.file_prefix}_OncoClini_Report.docx"
                st.sidebar.download_button("Download", word_data, file_name=name)








